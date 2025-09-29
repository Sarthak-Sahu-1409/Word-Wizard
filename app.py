
"""
The Word Wizard - A magical tool for managing PDF appendices in Word documents
Author: Expert Python Developer (Enhanced by Gemini)
Version: 3.2.0 (Robust Move Functionality)
"""

import os
import sys
import tempfile
import traceback
import threading
from pathlib import Path
from typing import List, Tuple, Optional, Set
import tkinter as tk
from tkinter import filedialog, messagebox

# --- Third-party imports with robust error handling ---
try:
    import customtkinter as ctk
except ImportError:
    messagebox.showerror("Dependency Error", "CustomTkinter not found. Please install it using: pip install customtkinter")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    messagebox.showerror("Dependency Error", "python-docx not found. Please install it using: pip install python-docx")
    sys.exit(1)

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    messagebox.showerror("Dependency Error", "pypdf not found. Please install it using: pip install pypdf")
    sys.exit(1)

try:
    from docx2pdf import convert
except ImportError:
    messagebox.showerror("Dependency Error", "docx2pdf not found. Please install it using: pip install docx2pdf")
    sys.exit(1)

# --- Windows-specific imports for live editing ---
IS_WINDOWS = sys.platform == "win32"
if IS_WINDOWS:
    try:
        import win32com.client as win32
        import pythoncom
    except ImportError:
        messagebox.showerror("Dependency Error", "pywin32 not found. Please install it using: pip install pywin32")
        sys.exit(1)

# --- Helper Function for COM Automation ---
def get_live_word_document():
    """
    Connects to a running instance of MS Word and returns the active document's full path.
    Returns None if Word is not running or no document is open.
    This function is Windows-only.
    """
    if not IS_WINDOWS:
        messagebox.showerror("Unsupported OS", "Live document detection is only available on Windows.")
        return None

    try:
        pythoncom.CoInitialize()
        word_app = win32.GetActiveObject("Word.Application")
        
        if word_app.Documents.Count > 0:
            doc = word_app.ActiveDocument
            return doc.FullName
        else:
            messagebox.showwarning("No Document", "Microsoft Word is running, but no document is open.")
            return None
            
    except pythoncom.com_error:
        messagebox.showerror("Word Not Found", "Microsoft Word is not running. Please open a document to use this feature.")
        return None
    except Exception as e:
        messagebox.showerror("Error", f"An unexpected error occurred while connecting to Word:\n{e}")
        return None
    finally:
        pythoncom.CoUninitialize()

# --- Main Application Class ---
class WordWizard(ctk.CTk):
    """
    Main application class for The Word Wizard.
    Provides a responsive GUI for adding PDF appendices to Word documents.
    """

    def __init__(self):
        super().__init__()

        self.title("The Word Wizard - PDF Appendix Manager")
        self.geometry("850x650")
        self.minsize(700, 550)

        ctk.set_appearance_mode("Dark")
        ctk.set_default_color_theme("dark-blue")

        self.word_doc_path: Optional[str] = None
        # (default_label, custom_title, path, page_count, page_range_str)
        self.appendix_list: List[Tuple[str, str, str, int, Optional[str]]] = []

        self._configure_grid()
        self._create_widgets()

        self.protocol("WM_DELETE_WINDOW", self._on_closing)

    def _configure_grid(self):
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

    def _create_widgets(self):
        main_frame = ctk.CTkFrame(self)
        main_frame.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")
        main_frame.grid_rowconfigure(2, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)

        header_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        header_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")
        
        ctk.CTkLabel(header_frame, text="🧙 The Word Wizard", font=ctk.CTkFont(size=28, weight="bold")).pack()
        ctk.CTkLabel(header_frame, text="Magically append PDFs to your Word documents", font=ctk.CTkFont(size=14), text_color=("gray60", "gray40")).pack()

        doc_frame = ctk.CTkFrame(main_frame)
        doc_frame.grid(row=1, column=0, padx=20, pady=10, sticky="ew")
        doc_frame.grid_columnconfigure(0, weight=1)
        
        ctk.CTkLabel(doc_frame, text="Step 1: Connect to Your Document", font=ctk.CTkFont(size=16, weight="bold")).grid(row=0, column=0, columnspan=2, padx=15, pady=(10, 5), sticky="w")
        
        self.connect_btn = ctk.CTkButton(doc_frame, text="🔗 Connect to Live Word Doc", command=self._connect_to_live_document)
        if not IS_WINDOWS:
            self.connect_btn.configure(state="disabled", text="Live Connect (Windows Only)")
        self.connect_btn.grid(row=1, column=0, padx=15, pady=10, sticky="w")
        
        self.doc_path_label = ctk.CTkLabel(doc_frame, text="Not connected.", text_color=("gray60", "gray40"), wraplength=400, justify="left")
        self.doc_path_label.grid(row=1, column=1, padx=15, pady=10, sticky="w")

        appendix_frame = ctk.CTkFrame(main_frame)
        appendix_frame.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")
        appendix_frame.grid_rowconfigure(1, weight=1)
        appendix_frame.grid_columnconfigure(0, weight=1)
        appendix_frame.grid_columnconfigure(1, weight=0)

        ctk.CTkLabel(appendix_frame, text="Step 2: Add and Organize PDF Appendices", font=ctk.CTkFont(size=16, weight="bold")).grid(row=0, column=0, columnspan=2, padx=15, pady=(10, 5), sticky="w")

        listbox_frame = ctk.CTkFrame(appendix_frame)
        listbox_frame.grid(row=1, column=0, padx=(15,0), pady=10, sticky="nsew")
        listbox_frame.grid_rowconfigure(0, weight=1)
        listbox_frame.grid_columnconfigure(0, weight=1)

        self.appendix_listbox = tk.Listbox(listbox_frame, font=("Segoe UI", 12), bg="#2E2E2E", fg="white", selectbackground="#1F6AA5", selectforeground="white", borderwidth=0, highlightthickness=0, activestyle="none")
        self.appendix_listbox.grid(row=0, column=0, sticky="nsew")

        listbox_scrollbar = ctk.CTkScrollbar(listbox_frame, command=self.appendix_listbox.yview)
        listbox_scrollbar.grid(row=0, column=1, sticky="ns")
        self.appendix_listbox.configure(yscrollcommand=listbox_scrollbar.set)
        
        self.appendix_listbox.bind("<Double-1>", self._rename_appendix)
        self.appendix_listbox.bind("<<ListboxSelect>>", self._update_control_button_states)

        controls_frame = ctk.CTkScrollableFrame(appendix_frame, label_text="Actions", width=150)
        controls_frame.grid(row=1, column=1, padx=15, pady=10, sticky="ns")
        
        self.add_pdf_btn = ctk.CTkButton(controls_frame, text="➕ Add", command=self._add_pdfs)
        self.add_pdf_btn.pack(pady=5, padx=5, fill="x")
        self.remove_btn = ctk.CTkButton(controls_frame, text="🗑️ Remove", command=self._remove_selected, state="disabled")
        self.remove_btn.pack(pady=5, padx=5, fill="x")
        
        self.configure_btn = ctk.CTkButton(controls_frame, text="⚙️ Configure Pages", command=self._configure_pages, state="disabled")
        self.configure_btn.pack(pady=(15, 5), padx=5, fill="x")

        self.move_up_btn = ctk.CTkButton(controls_frame, text="⬆️ Move Up", command=self._move_up, state="disabled")
        self.move_up_btn.pack(pady=5, padx=5, fill="x")
        self.move_down_btn = ctk.CTkButton(controls_frame, text="⬇️ Move Down", command=self._move_down, state="disabled")
        self.move_down_btn.pack(pady=5, padx=5, fill="x")
        
        action_frame = ctk.CTkFrame(main_frame)
        action_frame.grid(row=3, column=0, padx=20, pady=(10, 20), sticky="ew")
        action_frame.grid_columnconfigure(0, weight=1)

        self.generate_btn = ctk.CTkButton(action_frame, text="✨ Cast Spell & Generate Document", command=self._start_generation_thread, font=ctk.CTkFont(size=18, weight="bold"), height=50, state="disabled")
        self.generate_btn.grid(row=0, column=0, padx=15, pady=15, sticky="ew")
        
        self.progress_bar = ctk.CTkProgressBar(action_frame, mode='indeterminate')
        
        self.status_label = ctk.CTkLabel(action_frame, text="Ready for some magic...", font=ctk.CTkFont(size=12))
        self.status_label.grid(row=2, column=0, padx=15, pady=(0, 10), sticky="w")
    
    def _update_status(self, message: str, color: Optional[str] = None):
        def update():
            self.status_label.configure(text=message, text_color=color)
        self.after(0, update)

    def _connect_to_live_document(self):
        """Detects the active Word document and updates the UI."""
        file_path = get_live_word_document()
        if file_path:
            self.word_doc_path = file_path
            display_name = os.path.basename(file_path)
            self.doc_path_label.configure(text=f"Connected to: {display_name}")
            self._update_generate_button_state()

    def _add_pdfs(self):
        file_paths = filedialog.askopenfilenames(title="Select PDF Appendices", filetypes=[("PDF Files", "*.pdf")])
        if not file_paths: return
        for path in file_paths:
            try:
                with open(path, 'rb') as f:
                    reader = PdfReader(f)
                    page_count = len(reader.pages)
                default_label = f"Appendix {chr(65 + len(self.appendix_list))}"
                self.appendix_list.append((default_label, default_label, path, page_count, None))
            except Exception as e:
                messagebox.showerror("PDF Error", f"Could not read '{os.path.basename(path)}':\n{e}")
        self._refresh_appendix_listbox()
        # Set selection to the last added item
        self.appendix_listbox.selection_set(tk.END)
        self._update_control_button_states()
        self._update_generate_button_state()
        self._update_status(f"Added {len(file_paths)} PDF(s). Ready to organize.", "cyan")

    def _remove_selected(self):
        selected_indices = self.appendix_listbox.curselection()
        if not selected_indices: return
        
        # Store index to select after deletion
        select_after = selected_indices[0] - 1 if selected_indices[0] > 0 else 0
        
        for index in sorted(selected_indices, reverse=True):
            del self.appendix_list[index]
        
        self._refresh_appendix_listbox()
        
        # Reselect appropriately
        if self.appendix_list:
             self.appendix_listbox.selection_set(min(select_after, len(self.appendix_list)-1))

        self._update_generate_button_state()
        self._update_control_button_states()
        self._update_status("Appendix removed.", "yellow")

    def _move_item(self, direction: int):
        """Robustly moves a selected item up or down in the list."""
        selected_indices = self.appendix_listbox.curselection()
        # Ensure exactly one item is selected
        if not selected_indices or len(selected_indices) > 1:
            return
            
        idx = selected_indices[0]
        new_idx = idx + direction

        # Check if the new position is valid
        if 0 <= new_idx < len(self.appendix_list):
            # Move the item in the data list
            item = self.appendix_list.pop(idx)
            self.appendix_list.insert(new_idx, item)
            
            # Refresh the visual listbox from the data
            self._refresh_appendix_listbox()
            
            # Set selection and focus to the item's new position
            self.appendix_listbox.selection_set(new_idx)
            self.appendix_listbox.activate(new_idx)
            self.appendix_listbox.see(new_idx) # Ensures the moved item is visible

    def _move_up(self): self._move_item(-1)
    def _move_down(self): self._move_item(1)

    def _refresh_appendix_listbox(self):
        """Updates the listbox display and re-sequences the default Appendix labels."""
        # This function no longer manages selection; the calling function does.
        self.appendix_listbox.delete(0, tk.END)
        
        updated_list = []
        for i, (_, custom_title, pdf_path, page_count, page_range) in enumerate(self.appendix_list):
            new_default_label = f"Appendix {chr(65 + i)}"
            updated_list.append((new_default_label, custom_title, pdf_path, page_count, page_range))
            
            filename = os.path.basename(pdf_path)
            page_info = f"Pages: {page_range}" if page_range else f"{page_count} pages"
            display_text = f"{custom_title}  |  {page_info}  ({filename})"
            self.appendix_listbox.insert(tk.END, display_text)

        self.appendix_list = updated_list
        
    def _rename_appendix(self, event=None):
        """Handles the double-click event to rename an appendix."""
        selected_indices = self.appendix_listbox.curselection()
        if not selected_indices: return
        idx = selected_indices[0]
        
        default_label, current_title, path, count, page_range = self.appendix_list[idx]

        dialog = ctk.CTkInputDialog(
            title="Rename Appendix",
            text="Enter the new appendix title:"
        )
        dialog.entry.insert(0, current_title)
        
        new_title = dialog.get_input()

        if new_title and new_title.strip():
            self.appendix_list[idx] = (default_label, new_title, path, count, page_range)
            self._refresh_appendix_listbox()
            self.appendix_listbox.selection_set(idx)
            self._update_status(f"Renamed appendix to '{new_title}'.", "cyan")

    def _configure_pages(self):
        """Opens a dialog to configure the page range for the selected PDF."""
        selected_indices = self.appendix_listbox.curselection()
        if not selected_indices or len(selected_indices) > 1: return
        idx = selected_indices[0]
        
        _d_label, _c_title, _path, page_count, current_range = self.appendix_list[idx]

        dialog = ctk.CTkInputDialog(
            title="Configure Page Range",
            text=f'Enter page range (e.g., "1-5, 8, 10-12").\nTotal pages: {page_count}. Leave blank for all pages.'
        )
        if current_range:
            dialog.entry.insert(0, current_range)
            
        new_range_str = dialog.get_input()

        if new_range_str is not None:
            new_range_str = new_range_str.strip()
            if new_range_str and self._parse_page_range(new_range_str, page_count) is None:
                return
            
            final_range = new_range_str if new_range_str else None
            self.appendix_list[idx] = (*self.appendix_list[idx][:4], final_range)
            self._refresh_appendix_listbox()
            self.appendix_listbox.selection_set(idx)
            self._update_status("Updated page range.", "cyan")

    def _parse_page_range(self, range_str: str, max_pages: int) -> Optional[List[int]]:
        """Parses a page range string (e.g., "1-3, 5") into a list of 0-indexed pages."""
        if not range_str:
            return []
        
        pages_to_include: Set[int] = set()
        try:
            parts = range_str.split(',')
            for part in parts:
                part = part.strip()
                if not part: continue
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    if not (1 <= start <= end <= max_pages):
                        raise ValueError(f"Range '{part}' is invalid for a document with {max_pages} pages.")
                    pages_to_include.update(range(start, end + 1))
                else:
                    page = int(part)
                    if not (1 <= page <= max_pages):
                        raise ValueError(f"Page '{page}' is invalid for a document with {max_pages} pages.")
                    pages_to_include.add(page)
            return sorted([p - 1 for p in pages_to_include])
        except ValueError as e:
            messagebox.showerror("Invalid Page Range", f"The page range you entered is invalid.\n\nError: {e}\nPlease use formats like '1-5, 8, 10-12'.")
            return None

    def _update_generate_button_state(self):
        state = "normal" if self.word_doc_path and self.appendix_list else "disabled"
        self.generate_btn.configure(state=state)

    def _update_control_button_states(self, event=None):
        """Updates the state of Remove, Configure, and Move buttons based on selection."""
        selected_indices = self.appendix_listbox.curselection()
        
        if not selected_indices:
            self.remove_btn.configure(state="disabled")
            self.configure_btn.configure(state="disabled")
            self.move_up_btn.configure(state="disabled")
            self.move_down_btn.configure(state="disabled")
        elif len(selected_indices) == 1:
            self.remove_btn.configure(state="normal")
            self.configure_btn.configure(state="normal")
            self.move_up_btn.configure(state="normal")
            self.move_down_btn.configure(state="normal")
        else: 
            self.remove_btn.configure(state="normal")
            self.configure_btn.configure(state="disabled")
            self.move_up_btn.configure(state="disabled")
            self.move_down_btn.configure(state="disabled")

    def _set_controls_enabled(self, enabled: bool):
        def update():
            state = "normal" if enabled else "disabled"
            if IS_WINDOWS: self.connect_btn.configure(state=state)
            self.add_pdf_btn.configure(state=state)
            if enabled:
                self._update_control_button_states()
                self._update_generate_button_state()
            else:
                self.remove_btn.configure(state="disabled")
                self.configure_btn.configure(state="disabled")
                self.move_up_btn.configure(state="disabled")
                self.move_down_btn.configure(state="disabled")
                self.generate_btn.configure(state="disabled")
        self.after(0, update)

    def _start_generation_thread(self):
        if not self._validate_inputs(): return
        self._set_controls_enabled(False)
        self.progress_bar.grid(row=1, column=0, padx=15, pady=(5,10), sticky="ew")
        self.progress_bar.start()
        self._update_status("The magic is starting...", "cyan")
        thread = threading.Thread(target=self._generate_document_worker)
        thread.daemon = True
        thread.start()

    def _generate_document_worker(self):
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                self._update_status("Step 1/3: Converting main document to PDF...", "cyan")
                base_pdf_path = os.path.join(temp_dir, "base_document.pdf")
                self._convert_to_pdf(self.word_doc_path, base_pdf_path)

                self._update_status("Step 2/3: Generating appendix title pages...", "cyan")
                heading_pdf_paths = self._create_heading_pdfs(temp_dir)
                
                self._update_status("Step 3/3: Merging all documents...", "cyan")
                final_path = self._merge_and_save_final_pdf(base_pdf_path, heading_pdf_paths)

                if final_path:
                    self._update_status(f"🎉 Success! Document saved to {os.path.basename(final_path)}", "lime green")
                    self.after(0, lambda: messagebox.showinfo("Success", f"Your magical document has been created!\n\nSaved at: {final_path}"))
                else:
                    self.update()
                    self._update_status("Process cancelled by user.", "yellow")

        except Exception as e:
            error_details = traceback.format_exc()
            self._update_status(f"❌ An error occurred: {e}", "red")
            self.after(0, lambda: messagebox.showerror("Generation Error", f"An unexpected error occurred:\n\n{e}\n\nDetails:\n{error_details}"))
        finally:
            def finalize_ui():
                self.progress_bar.stop()
                self.progress_bar.grid_forget()
                self._set_controls_enabled(True)
            self.after(0, finalize_ui)

    def _validate_inputs(self) -> bool:
        if not self.word_doc_path or not os.path.exists(self.word_doc_path):
            messagebox.showwarning("Validation Error", "Please connect to a valid Word document.")
            return False
        if not self.appendix_list:
            messagebox.showwarning("Validation Error", "Please add at least one PDF appendix.")
            return False
        for _, _, pdf_path, _, _ in self.appendix_list:
            if not os.path.exists(pdf_path):
                messagebox.showerror("File Not Found", f"The PDF file could not be found:\n{pdf_path}")
                return False
        return True

    def _create_heading_pdfs(self, temp_dir: str) -> List[str]:
        """Creates a professional, centered title page for each appendix."""
        heading_paths = []
        for i, (_d_label, custom_title, pdf_path, _, _) in enumerate(self.appendix_list):
            doc = Document()
            
            for _ in range(8):
                doc.add_paragraph()

            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run(custom_title)
            font_title = run_title.font
            font_title.name = 'Calibri'
            font_title.size = Pt(36)
            font_title.bold = True

            doc.add_paragraph() 
            
            p_subtitle = doc.add_paragraph()
            p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_subtitle = p_subtitle.add_run(f"({os.path.basename(pdf_path)})")
            font_subtitle = run_subtitle.font
            font_subtitle.name = 'Calibri'
            font_subtitle.size = Pt(12)
            font_subtitle.italic = True
            
            temp_docx_path = os.path.join(temp_dir, f"heading_{i}.docx")
            temp_pdf_path = os.path.join(temp_dir, f"heading_{i}.pdf")
            doc.save(temp_docx_path)
            self._convert_to_pdf(temp_docx_path, temp_pdf_path)
            heading_paths.append(temp_pdf_path)
        return heading_paths

    def _convert_to_pdf(self, input_docx: str, output_pdf: str):
        try:
            convert(input_docx, output_pdf)
            if not os.path.exists(output_pdf):
                raise RuntimeError("Conversion failed silently. Output PDF not found.")
        except Exception as e:
            if IS_WINDOWS and ("com_error" in str(e).lower() or "word" in str(e).lower()):
                 raise RuntimeError("PDF conversion failed. Ensure MS Word is installed and not busy.")
            elif not IS_WINDOWS and "libreoffice" in str(e).lower():
                 raise RuntimeError("PDF conversion failed. Ensure LibreOffice is installed and accessible.")
            else:
                 raise e

    def _merge_and_save_final_pdf(self, base_pdf_path: str, heading_pdf_paths: List[str]) -> Optional[str]:
        save_path_container = []
        def prompt_save():
            initial_filename = f"{Path(self.word_doc_path).stem}_with_appendices.pdf"
            path = filedialog.asksaveasfilename(title="Save Final Document", initialfile=initial_filename, defaultextension=".pdf", filetypes=[("PDF Documents", "*.pdf")])
            save_path_container.append(path)
        
        self.after(0, prompt_save)
        while not save_path_container:
            import time
            time.sleep(0.1)
        
        output_path = save_path_container[0]
        if not output_path: return None

        writer = PdfWriter()
        writer.append(base_pdf_path)
        for i, (_, _, appendix_path, page_count, page_range_str) in enumerate(self.appendix_list):
            writer.append(heading_pdf_paths[i])
            
            pages_to_add = self._parse_page_range(page_range_str, page_count)
            
            if pages_to_add is None and page_range_str:
                raise RuntimeError(f"Could not proceed due to invalid page range for {os.path.basename(appendix_path)}")

            if page_range_str and pages_to_add is not None:
                reader = PdfReader(appendix_path)
                for page_num in pages_to_add:
                    if page_num < len(reader.pages):
                        writer.add_page(reader.pages[page_num])
            else:
                 writer.append(appendix_path)
                  
        with open(output_path, "wb") as f_out:
            writer.write(f_out)
        
        return output_path

    def _on_closing(self):
        if messagebox.askokcancel("Quit", "Do you want to exit The Word Wizard?"):
            self.destroy()

if __name__ == "__main__":
    app = WordWizard()
    app.mainloop()